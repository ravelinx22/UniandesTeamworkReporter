from docx import Document
from docx.shared import Inches
from pptx.chart.data import ChartData
from pptx.enum.chart import (
    XL_CHART_TYPE, XL_LEGEND_POSITION, XL_LABEL_POSITION)
from docx.shared import Pt
from docx.shared import RGBColor

########################################################
# CLASSES ##############################################
########################################################


class Student:
    def __init__(self, name):
        self.name = name.title()

    def setAsignadas(self, asignadas):
        self.asignadas = asignadas

    def setCompletadas(self, completadas):
        self.completadas = completadas

    def getCompletitud(self):
        return float(self.asignadas / self.completadas) * 100

########################################################
# HELPERS ##############################################
########################################################


if hasattr(__builtins__, 'raw_input'):
    input = raw_input


def toPercentageArray(total, number_array):
    percentage_array = []
    for number in number_array:
        if(int(total) != 0):
            percentage_array.append(float(number) / float(total))
        else:
            percentage_array.append(0)
    return percentage_array


def toNameArray(members_list):
    names = []
    for member in members_list:
        if not(isinstance(member, str)):
            names.append(member.name)
        else:
            names.append(member)
    return names

########################################################
# CREATE REPORT ########################################
########################################################


def create_pie_chart(document_to_add_chart, chart_name, categories, series):
    chart_data = ChartData()
    chart_data.categories = categories
    chart_data.add_series(chart_name, series)

    x, y, cx, cy = Inches(2), Inches(2), Inches(6), Inches(3.5)

    chart = document.add_chart(XL_CHART_TYPE.PIE, x, y, cx, cy, chart_data)

    chart.has_legend = True
    chart.legend.position = XL_LEGEND_POSITION.BOTTOM
    chart.legend.include_in_layout = False

    chart.plots[0].has_data_labels = True
    data_labels = chart.plots[0].data_labels
    data_labels.number_format = '0.00%'
    data_labels.position = XL_LABEL_POSITION.OUTSIDE_END

########################################################
# MAIN #################################################
########################################################

# Data Setup


global group_name
global members

testGroup = [Student("William Ravelo M")]

otherGroup = [Student("Ricardo Ravelo M")]


def escoger_grupo():
    option = -1

    # TODO: Change number of options. The number of options represents the number of groups in the course
    while(option <= 0 or option > 2):
        print("Escoga que grupo a hacer reporte: ")
        # TODO: Add an option for every group
        print("1. Test Group")
        print("2. Other Group")
        option = int(input('Ingrese el numero del grupo: '))

    global members
    global group_name

    # TODO. Add an option if for every group.
    if(option == 1):
        members = testGroup[:]
        group_name = "Test Group"
    else:
        members = otherGroup[:]
        group_name = "Other Group"

########################################################
# GET DATA #############################################
########################################################

cycle = input("Numero del ciclo: ")
cycle_week = input("Numero de semana: ")

escoger_grupo()

total_tasks = input("Numero de tareas totales: ")

# Document Setup
document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10)

# Header
header = document.add_paragraph().add_run("Ciclo " + cycle
                                          + " Semana " + cycle_week)
header.bold = True
font = header.font
font.size = Pt(20)
font.color.rgb = RGBColor(0, 0, 0)


paragraph = document.add_paragraph()
paragraph.add_run('Nombre Grupo:\n').bold = True
paragraph.add_run(group_name)

paragraph = document.add_paragraph()
paragraph.add_run('Integrantes:\n').bold = True

for member in members:
    paragraph.add_run(member.name + '\n')

########################################################
# BASIC VERSION ########################################
########################################################

# Responsabilidades
categories = ['Tareas asignadas', 'Tareas sin  asignar']

assigned = input("Tareas asignadas: ")
not_assigned = input("Tareas sin asignar: ")

series = [float(assigned), float(not_assigned)]
create_pie_chart(document, 'Responsabilidades',  categories,
                 toPercentageArray(total_tasks, series))

# Tiempos
categories = ['Tareas con tiempos', 'Tareas sin tiempos']

with_time = input("Tareas con tiempos: ")
without_time = input("Tareas sin tiempos: ")

series = [float(with_time), float(without_time)]
create_pie_chart(document, 'Tiempos',  categories,
                 toPercentageArray(total_tasks, series))

# Tareas completadas
categories = ['Completadas A Tiempo', 'Completadas Tarde', 'Sin Completar']

time = input("Completadas A Tiempo: ")
late = input("Completadas Tarde: ")
not_completed = input("Sin Completar: ")

series = [float(time), float(late), float(not_completed)]
create_pie_chart(document, 'Tareas completadas',  categories,
                 toPercentageArray(total_tasks, series))

# Tareas asignadas por integrante
series = []
for member in members:
    assigned_by_member = input("Asignadas a " + member.name + ": ")
    member.setAsignadas(float(assigned_by_member))
    series.append(float(assigned_by_member))

member_without_assign = input("Sin asignar: ")
series.append(float(member_without_assign))

members_with_without_assign = members[:]
members_with_without_assign.append("Sin asignar")

create_pie_chart(document, 'Tareas asignadas por integrante',
                 toNameArray(members_with_without_assign),
                 toPercentageArray(total_tasks, series))

# Tareas completadas por integrante
series = []
for member in members:
    complete_by_member = input("Completadas por " + member.name + ": ")
    member.setCompletadas(float(complete_by_member))
    series.append(float(complete_by_member))

member_without_complete = input("Sin completar: ")
series.append(float(member_without_complete))

members_with_without_complete = members[:]
members_with_without_complete.append("Sin completar")

create_pie_chart(document, 'Tareas completadas por integrante',
                 toNameArray(members_with_without_complete),
                 toPercentageArray(total_tasks, series))

########################################################
# PROFESSOR VERSION ####################################
########################################################

print(members)

# Save
document.save(group_name + '.docx')
